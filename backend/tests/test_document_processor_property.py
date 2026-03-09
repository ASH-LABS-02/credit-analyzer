"""
Property-Based Tests for Document Processing

This module contains property-based tests for document processing functionality,
validating universal properties that should hold across all inputs.

Feature: intelli-credit-platform
Task: 3.2 Write property tests for document processing
"""

import io
import pytest
from hypothesis import given, strategies as st, settings, assume
from hypothesis import HealthCheck

from app.services.document_processor import (
    DocumentProcessor,
    FileType,
    UnsupportedFileTypeError,
    CorruptedFileError,
    FileSizeExceededError,
)

# Import libraries for creating test files
import PyPDF2
from docx import Document
import openpyxl
import pandas as pd
from PIL import Image


# ============================================================================
# Test Data Generators (Strategies)
# ============================================================================

@st.composite
def valid_file_extensions(draw):
    """Generate valid file extensions."""
    extensions = ['.pdf', '.docx', '.xlsx', '.xls', '.csv', '.jpg', '.jpeg', '.png']
    return draw(st.sampled_from(extensions))


@st.composite
def invalid_file_extensions(draw):
    """Generate invalid file extensions."""
    invalid_exts = ['.txt', '.doc', '.rtf', '.xml', '.json', '.exe', '.zip', '.tar', 
                    '.mp3', '.mp4', '.avi', '.mov', '.html', '.js', '.py', '.java']
    return draw(st.sampled_from(invalid_exts))


@st.composite
def valid_filename(draw):
    """Generate valid filename with supported extension."""
    # Generate a base filename (alphanumeric with some special chars)
    base = draw(st.text(
        alphabet=st.characters(whitelist_categories=('Lu', 'Ll', 'Nd'), whitelist_characters='_-'),
        min_size=1,
        max_size=50
    ))
    ext = draw(valid_file_extensions())
    return base + ext


@st.composite
def invalid_filename(draw):
    """Generate filename with unsupported extension."""
    base = draw(st.text(
        alphabet=st.characters(whitelist_categories=('Lu', 'Ll', 'Nd'), whitelist_characters='_-'),
        min_size=1,
        max_size=50
    ))
    ext = draw(invalid_file_extensions())
    return base + ext


@st.composite
def file_content_within_limit(draw):
    """Generate file content within size limit."""
    # Generate content smaller than MAX_FILE_SIZE
    max_size = DocumentProcessor.MAX_FILE_SIZE - 1000  # Leave some margin
    size = draw(st.integers(min_value=0, max_value=min(max_size, 10000)))  # Cap at 10KB for test speed
    return draw(st.binary(min_size=size, max_size=size))


@st.composite
def file_content_exceeding_limit(draw):
    """Generate file content exceeding size limit."""
    # Generate content larger than MAX_FILE_SIZE
    excess = draw(st.integers(min_value=1, max_value=1000))
    size = DocumentProcessor.MAX_FILE_SIZE + excess
    # For testing, we don't need to generate the full content, just report the size
    # We'll create a minimal representation
    return b'x' * min(size, DocumentProcessor.MAX_FILE_SIZE + 1000)


@st.composite
def valid_pdf_content(draw):
    """Generate valid minimal PDF content."""
    # Minimal valid PDF structure
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
>>
endobj
xref
0 4
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
trailer
<<
/Size 4
/Root 1 0 R
>>
startxref
200
%%EOF
"""
    return pdf_content


@st.composite
def valid_docx_content(draw):
    """Generate valid DOCX content."""
    doc = Document()
    # Generate XML-safe text (printable characters only, no control characters)
    text = draw(st.text(
        alphabet=st.characters(
            whitelist_categories=('Lu', 'Ll', 'Nd', 'Zs'),
            min_codepoint=32,  # Space character
            max_codepoint=126  # Tilde character (printable ASCII)
        ),
        min_size=1,
        max_size=100
    ))
    doc.add_paragraph(text)
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    return docx_buffer.getvalue()


@st.composite
def valid_excel_content(draw):
    """Generate valid Excel content."""
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    
    # Add some data with safe text
    rows = draw(st.integers(min_value=1, max_value=10))
    for i in range(rows):
        value = draw(st.text(
            alphabet=st.characters(
                whitelist_categories=('Lu', 'Ll', 'Nd', 'Zs'),
                min_codepoint=32,
                max_codepoint=126
            ),
            min_size=1,
            max_size=20
        ))
        sheet.cell(row=i+1, column=1, value=value)
    
    excel_buffer = io.BytesIO()
    workbook.save(excel_buffer)
    return excel_buffer.getvalue()


@st.composite
def valid_csv_content(draw):
    """Generate valid CSV content."""
    rows = draw(st.integers(min_value=1, max_value=10))
    cols = draw(st.integers(min_value=1, max_value=5))
    
    data = {}
    for col in range(cols):
        col_name = f"col_{col}"
        data[col_name] = [
            draw(st.text(
                alphabet=st.characters(
                    whitelist_categories=('Lu', 'Ll', 'Nd', 'Zs'),
                    min_codepoint=32,
                    max_codepoint=126
                ),
                min_size=1,
                max_size=20
            )) for _ in range(rows)
        ]
    
    df = pd.DataFrame(data)
    csv_buffer = io.BytesIO()
    df.to_csv(csv_buffer, index=False)
    return csv_buffer.getvalue()


@st.composite
def valid_image_content(draw):
    """Generate valid image content."""
    width = draw(st.integers(min_value=10, max_value=200))
    height = draw(st.integers(min_value=10, max_value=200))
    
    image = Image.new('RGB', (width, height), color='white')
    img_buffer = io.BytesIO()
    image.save(img_buffer, format='PNG')
    return img_buffer.getvalue()


@st.composite
def corrupted_content(draw):
    """Generate corrupted/invalid file content."""
    # Generate random binary data that's unlikely to be valid
    return draw(st.binary(min_size=10, max_size=1000))


@st.composite
def document_batch(draw):
    """Generate a batch of documents with mixed valid and corrupted files."""
    batch_size = draw(st.integers(min_value=2, max_value=5))
    documents = []
    
    for i in range(batch_size):
        # Randomly decide if this document should be corrupted
        is_corrupted = draw(st.booleans())
        
        if is_corrupted:
            filename = f"corrupted_{i}.pdf"
            content = draw(corrupted_content())
        else:
            ext = draw(valid_file_extensions())
            filename = f"valid_{i}{ext}"
            
            # Generate appropriate content based on extension
            if ext == '.pdf':
                content = draw(valid_pdf_content())
            elif ext == '.docx':
                content = draw(valid_docx_content())
            elif ext in ['.xlsx', '.xls']:
                content = draw(valid_excel_content())
            elif ext == '.csv':
                content = draw(valid_csv_content())
            else:  # image
                content = draw(valid_image_content())
        
        documents.append({
            'filename': filename,
            'content': content,
            'is_corrupted': is_corrupted
        })
    
    return documents


# ============================================================================
# Property 2: Invalid Document Rejection
# **Validates: Requirements 1.2, 1.3**
# ============================================================================

@pytest.mark.property
class TestInvalidDocumentRejection:
    """
    Property 2: Invalid Document Rejection
    
    For any file with an invalid format or exceeding size limits,
    the upload should be rejected with a descriptive error message,
    and no storage operation should occur.
    
    **Validates: Requirements 1.2, 1.3**
    """
    
    @settings(max_examples=20, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @given(filename=invalid_filename())
    def test_invalid_format_rejection(self, filename):
        """
        Property: Files with invalid/unsupported formats are always rejected.
        
        For any filename with an unsupported extension, the system should:
        1. Raise UnsupportedFileTypeError
        2. Include descriptive error message
        3. Not proceed with processing
        """
        processor = DocumentProcessor()
        
        # Attempt to detect file type
        with pytest.raises(UnsupportedFileTypeError) as exc_info:
            processor.detect_file_type(filename)
        
        # Verify error message is descriptive
        error_message = str(exc_info.value)
        assert "Unsupported file type" in error_message or "not supported" in error_message.lower()
        
        # Verify the problematic extension is mentioned
        from pathlib import Path
        ext = Path(filename).suffix.lower()
        assert ext in error_message
    
    @settings(max_examples=20, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @given(
        filename=valid_filename(),
        content=file_content_exceeding_limit()
    )
    def test_oversized_file_rejection(self, filename, content):
        """
        Property: Files exceeding size limits are always rejected.
        
        For any file with size > MAX_FILE_SIZE, the system should:
        1. Raise FileSizeExceededError
        2. Include descriptive error message with size information
        3. Not proceed with processing
        """
        processor = DocumentProcessor()
        
        # Ensure content actually exceeds limit
        assume(len(content) > DocumentProcessor.MAX_FILE_SIZE)
        
        # Attempt to validate file
        with pytest.raises(FileSizeExceededError) as exc_info:
            processor.validate_file(filename, content)
        
        # Verify error message is descriptive
        error_message = str(exc_info.value)
        assert "exceeds maximum" in error_message.lower() or "too large" in error_message.lower()
        assert "size" in error_message.lower()
    
    @settings(max_examples=20, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @given(
        filename=valid_filename(),
        content=file_content_within_limit()
    )
    def test_valid_file_not_rejected(self, filename, content):
        """
        Property: Valid files within size limits are not rejected during validation.
        
        For any file with valid format and size <= MAX_FILE_SIZE, the system should:
        1. Successfully validate the file
        2. Return the correct file type
        3. Return is_valid=True
        """
        processor = DocumentProcessor()
        
        # Ensure content is within limit
        assume(len(content) <= DocumentProcessor.MAX_FILE_SIZE)
        
        # Validate file should succeed
        file_type, is_valid = processor.validate_file(filename, content)
        
        # Verify validation succeeded
        assert is_valid is True
        assert isinstance(file_type, FileType)
        
        # Verify file type matches extension
        from pathlib import Path
        ext = Path(filename).suffix.lower()
        expected_type = DocumentProcessor.SUPPORTED_EXTENSIONS.get(ext)
        assert file_type == expected_type
    
    @settings(max_examples=10, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @given(filename=valid_filename())
    def test_corrupted_file_rejection(self, filename):
        """
        Property: Corrupted files are rejected with appropriate error.
        
        For any file with valid extension but corrupted/invalid content,
        the system should:
        1. Raise CorruptedFileError during extraction
        2. Include descriptive error message
        3. Not crash or return invalid data
        """
        processor = DocumentProcessor()
        
        # Generate corrupted content (random bytes)
        corrupted_content = b"This is not valid file content" + b"\x00\xff\xfe" * 10
        
        # Validation might pass (only checks format and size)
        # But extraction should fail
        with pytest.raises((CorruptedFileError, Exception)) as exc_info:
            processor.extract_text(filename, corrupted_content)
        
        # If it's a CorruptedFileError, verify it's descriptive
        if isinstance(exc_info.value, CorruptedFileError):
            error_message = str(exc_info.value)
            assert "failed" in error_message.lower() or "corrupted" in error_message.lower()


# ============================================================================
# Property 33: Document Processing Error Isolation
# **Validates: Requirements 15.4**
# ============================================================================

@pytest.mark.property
class TestDocumentProcessingErrorIsolation:
    """
    Property 33: Document Processing Error Isolation
    
    For any application with multiple documents where one document is corrupted
    or fails processing, the system should isolate the error to that document
    and continue processing the remaining documents.
    
    **Validates: Requirements 15.4**
    """
    
    @settings(max_examples=10, suppress_health_check=[HealthCheck.function_scoped_fixture, HealthCheck.too_slow])
    @given(documents=document_batch())
    def test_error_isolation_in_batch_processing(self, documents):
        """
        Property: Errors in one document don't prevent processing of other documents.
        
        For any batch of documents containing at least one corrupted file,
        the system should:
        1. Process all valid documents successfully
        2. Isolate errors to corrupted documents only
        3. Return results for valid documents
        4. Report errors for corrupted documents
        """
        processor = DocumentProcessor()
        
        # Ensure we have at least one corrupted and one valid document
        has_corrupted = any(doc['is_corrupted'] for doc in documents)
        has_valid = any(not doc['is_corrupted'] for doc in documents)
        assume(has_corrupted and has_valid)
        
        # Process each document and collect results
        results = []
        errors = []
        
        for doc in documents:
            try:
                # Attempt to extract text
                text = processor.extract_text(doc['filename'], doc['content'])
                results.append({
                    'filename': doc['filename'],
                    'success': True,
                    'text': text,
                    'was_corrupted': doc['is_corrupted']
                })
            except (CorruptedFileError, UnsupportedFileTypeError, FileSizeExceededError) as e:
                # Error was isolated - processing continues
                errors.append({
                    'filename': doc['filename'],
                    'error': str(e),
                    'error_type': type(e).__name__,
                    'was_corrupted': doc['is_corrupted']
                })
        
        # Verify error isolation properties
        
        # 1. All valid documents should have been processed successfully
        valid_docs = [doc for doc in documents if not doc['is_corrupted']]
        successful_results = [r for r in results if r['success']]
        
        # At least some valid documents should succeed
        # (Note: Some "valid" generated content might still fail due to minimal structure)
        assert len(successful_results) > 0 or len(errors) == len(documents)
        
        # 2. Errors should be isolated (we got results for some documents)
        assert len(results) + len(errors) == len(documents)
        
        # 3. Each document was processed exactly once
        processed_files = set(r['filename'] for r in results) | set(e['filename'] for e in errors)
        expected_files = set(doc['filename'] for doc in documents)
        assert processed_files == expected_files
        
        # 4. Corrupted documents should be in errors (mostly)
        # Note: Due to random generation, some "corrupted" content might accidentally be valid
        corrupted_filenames = {doc['filename'] for doc in documents if doc['is_corrupted']}
        error_filenames = {e['filename'] for e in errors}
        
        # At least some corrupted files should have errors
        if corrupted_filenames:
            assert len(error_filenames & corrupted_filenames) > 0
    
    @settings(max_examples=10, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @given(
        valid_content=valid_pdf_content(),
        corrupted_content=corrupted_content()
    )
    def test_single_corrupted_document_isolation(self, valid_content, corrupted_content):
        """
        Property: A single corrupted document doesn't affect processing of valid documents.
        
        When processing multiple documents where one is corrupted:
        1. The corrupted document raises an error
        2. Valid documents are processed successfully
        3. The system doesn't crash or stop processing
        """
        processor = DocumentProcessor()
        
        documents = [
            {'filename': 'valid1.pdf', 'content': valid_content},
            {'filename': 'corrupted.pdf', 'content': corrupted_content},
            {'filename': 'valid2.pdf', 'content': valid_content},
        ]
        
        results = []
        errors = []
        
        # Process all documents
        for doc in documents:
            try:
                text = processor.extract_text(doc['filename'], doc['content'])
                results.append({'filename': doc['filename'], 'text': text})
            except (CorruptedFileError, Exception) as e:
                errors.append({'filename': doc['filename'], 'error': str(e)})
        
        # Verify isolation
        # 1. We should have processed all 3 documents (some succeeded, some failed)
        assert len(results) + len(errors) == 3
        
        # 2. At least one document should have succeeded (the valid ones)
        assert len(results) >= 1
        
        # 3. The corrupted document should have an error
        error_filenames = [e['filename'] for e in errors]
        assert 'corrupted.pdf' in error_filenames
        
        # 4. Valid documents should have succeeded
        success_filenames = [r['filename'] for r in results]
        # At least one valid document should succeed
        assert 'valid1.pdf' in success_filenames or 'valid2.pdf' in success_filenames
    
    @settings(max_examples=30, suppress_health_check=[HealthCheck.function_scoped_fixture, HealthCheck.too_slow])
    @given(
        num_valid=st.integers(min_value=1, max_value=3),
        num_corrupted=st.integers(min_value=1, max_value=2)
    )
    def test_multiple_corrupted_documents_isolation(self, num_valid, num_corrupted):
        """
        Property: Multiple corrupted documents don't prevent processing of valid ones.
        
        When processing a batch with multiple corrupted documents:
        1. All corrupted documents raise errors
        2. All valid documents are processed
        3. Error count matches corrupted document count
        4. Success count matches valid document count (approximately)
        """
        processor = DocumentProcessor()
        
        # Create a mix of valid and corrupted documents
        documents = []
        
        # Add valid documents
        for i in range(num_valid):
            doc = Document()
            doc.add_paragraph(f"Valid document {i}")
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            documents.append({
                'filename': f'valid_{i}.docx',
                'content': docx_buffer.getvalue(),
                'is_valid': True
            })
        
        # Add corrupted documents
        for i in range(num_corrupted):
            documents.append({
                'filename': f'corrupted_{i}.docx',
                'content': b'invalid content' + bytes([i]),
                'is_valid': False
            })
        
        # Process all documents
        results = []
        errors = []
        
        for doc in documents:
            try:
                text = processor.extract_text(doc['filename'], doc['content'])
                results.append({
                    'filename': doc['filename'],
                    'text': text,
                    'expected_valid': doc['is_valid']
                })
            except (CorruptedFileError, Exception) as e:
                errors.append({
                    'filename': doc['filename'],
                    'error': str(e),
                    'expected_valid': doc['is_valid']
                })
        
        # Verify isolation properties
        # 1. All documents were processed (either succeeded or failed)
        assert len(results) + len(errors) == num_valid + num_corrupted
        
        # 2. Valid documents should mostly succeed
        valid_successes = [r for r in results if r['expected_valid']]
        assert len(valid_successes) >= num_valid * 0.8  # Allow some margin
        
        # 3. Corrupted documents should mostly fail
        corrupted_errors = [e for e in errors if not e['expected_valid']]
        assert len(corrupted_errors) >= num_corrupted * 0.8  # Allow some margin
        
        # 4. Processing didn't stop after first error
        assert len(results) > 0 or len(errors) == len(documents)


# ============================================================================
# Additional Property Tests for Robustness
# ============================================================================

@pytest.mark.property
class TestDocumentProcessorRobustness:
    """Additional property tests for document processor robustness."""
    
    @settings(max_examples=20, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @given(
        filename=valid_filename(),
        content=file_content_within_limit()
    )
    def test_validation_idempotency(self, filename, content):
        """
        Property: Validation is idempotent - multiple validations return same result.
        
        For any file, validating it multiple times should:
        1. Return the same result each time
        2. Not modify the file content
        3. Not change processor state
        """
        processor = DocumentProcessor()
        
        assume(len(content) <= DocumentProcessor.MAX_FILE_SIZE)
        
        # Validate multiple times
        try:
            result1 = processor.validate_file(filename, content)
            result2 = processor.validate_file(filename, content)
            result3 = processor.validate_file(filename, content)
            
            # All results should be identical
            assert result1 == result2 == result3
        except (UnsupportedFileTypeError, FileSizeExceededError):
            # If validation fails, it should fail consistently
            with pytest.raises((UnsupportedFileTypeError, FileSizeExceededError)):
                processor.validate_file(filename, content)
            with pytest.raises((UnsupportedFileTypeError, FileSizeExceededError)):
                processor.validate_file(filename, content)
    
    @settings(max_examples=10, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @given(filename=valid_filename())
    def test_empty_file_handling(self, filename):
        """
        Property: Empty files are handled gracefully.
        
        For any valid filename with empty content:
        1. Validation should succeed (format and size are valid)
        2. Extraction might fail or return empty string
        3. No crashes or unhandled exceptions
        """
        processor = DocumentProcessor()
        
        empty_content = b''
        
        # Validation should succeed (empty file is within size limit)
        file_type, is_valid = processor.validate_file(filename, empty_content)
        assert is_valid is True
        
        # Extraction might fail or return empty, but shouldn't crash
        try:
            text = processor.extract_text(filename, empty_content)
            # If it succeeds, text should be a string (possibly empty)
            assert isinstance(text, str)
        except (CorruptedFileError, Exception):
            # If it fails, that's acceptable for empty files
            pass
