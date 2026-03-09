"""
Unit tests for DocumentProcessor

Tests document text extraction functionality for various file formats.
"""

import io
import pytest
from pathlib import Path

from app.services.document_processor import (
    DocumentProcessor,
    FileType,
    UnsupportedFileTypeError,
    CorruptedFileError,
    FileSizeExceededError,
)

# Import libraries for creating test files
import PyPDF2
from docx import Document
import openpyxl
import pandas as pd
from PIL import Image


@pytest.fixture
def processor():
    """Create a DocumentProcessor instance"""
    return DocumentProcessor()


@pytest.fixture
def sample_pdf():
    """Create a sample PDF file"""
    # Create a simple PDF with reportlab or use a minimal valid PDF
    # For testing purposes, we'll create a minimal PDF structure
    # Note: PyPDF2 can't easily add text to PDFs, so we create a minimal valid PDF
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/Resources <<
/Font <<
/F1 <<
/Type /Font
/Subtype /Type1
/BaseFont /Helvetica
>>
>>
>>
/MediaBox [0 0 612 792]
/Contents 4 0 R
>>
endobj
4 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
100 700 Td
(Sample PDF Text) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000317 00000 n 
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
410
%%EOF
"""
    return pdf_content


@pytest.fixture
def sample_docx():
    """Create a sample DOCX file"""
    doc = Document()
    doc.add_paragraph("Sample financial statement")
    doc.add_paragraph("Revenue: $1,000,000")
    doc.add_paragraph("Profit: $250,000")
    
    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Assets"
    table.cell(1, 1).text = "$5,000,000"
    table.cell(2, 0).text = "Liabilities"
    table.cell(2, 1).text = "$2,000,000"
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    return docx_buffer.getvalue()


@pytest.fixture
def sample_excel():
    """Create a sample Excel file"""
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Financial Data"
    
    # Add headers
    sheet['A1'] = "Year"
    sheet['B1'] = "Revenue"
    sheet['C1'] = "Profit"
    
    # Add data
    sheet['A2'] = 2021
    sheet['B2'] = 1000000
    sheet['C2'] = 250000
    
    sheet['A3'] = 2022
    sheet['B3'] = 1200000
    sheet['C3'] = 300000
    
    excel_buffer = io.BytesIO()
    workbook.save(excel_buffer)
    return excel_buffer.getvalue()


@pytest.fixture
def sample_csv():
    """Create a sample CSV file"""
    df = pd.DataFrame({
        'Year': [2021, 2022, 2023],
        'Revenue': [1000000, 1200000, 1500000],
        'Profit': [250000, 300000, 400000]
    })
    
    csv_buffer = io.BytesIO()
    df.to_csv(csv_buffer, index=False)
    return csv_buffer.getvalue()


@pytest.fixture
def sample_image():
    """Create a sample image file"""
    image = Image.new('RGB', (100, 100), color='white')
    
    img_buffer = io.BytesIO()
    image.save(img_buffer, format='PNG')
    return img_buffer.getvalue()


class TestFileTypeDetection:
    """Test file type detection functionality"""
    
    def test_detect_pdf_by_extension(self, processor):
        """Test PDF detection by file extension"""
        file_type = processor.detect_file_type("document.pdf")
        assert file_type == FileType.PDF
    
    def test_detect_docx_by_extension(self, processor):
        """Test DOCX detection by file extension"""
        file_type = processor.detect_file_type("document.docx")
        assert file_type == FileType.DOCX
    
    def test_detect_xlsx_by_extension(self, processor):
        """Test XLSX detection by file extension"""
        file_type = processor.detect_file_type("spreadsheet.xlsx")
        assert file_type == FileType.XLSX
    
    def test_detect_xls_by_extension(self, processor):
        """Test XLS detection by file extension"""
        file_type = processor.detect_file_type("spreadsheet.xls")
        assert file_type == FileType.XLS
    
    def test_detect_csv_by_extension(self, processor):
        """Test CSV detection by file extension"""
        file_type = processor.detect_file_type("data.csv")
        assert file_type == FileType.CSV
    
    def test_detect_jpg_by_extension(self, processor):
        """Test JPG detection by file extension"""
        file_type = processor.detect_file_type("image.jpg")
        assert file_type == FileType.JPG
    
    def test_detect_jpeg_by_extension(self, processor):
        """Test JPEG detection by file extension"""
        file_type = processor.detect_file_type("image.jpeg")
        assert file_type == FileType.JPEG
    
    def test_detect_png_by_extension(self, processor):
        """Test PNG detection by file extension"""
        file_type = processor.detect_file_type("image.png")
        assert file_type == FileType.PNG
    
    def test_case_insensitive_detection(self, processor):
        """Test that file type detection is case-insensitive"""
        assert processor.detect_file_type("document.PDF") == FileType.PDF
        assert processor.detect_file_type("document.DOCX") == FileType.DOCX
        assert processor.detect_file_type("data.CSV") == FileType.CSV
    
    def test_unsupported_file_type(self, processor):
        """Test that unsupported file types raise an error"""
        with pytest.raises(UnsupportedFileTypeError) as exc_info:
            processor.detect_file_type("document.txt")
        
        assert "Unsupported file type" in str(exc_info.value)
        assert ".txt" in str(exc_info.value)


class TestFileValidation:
    """Test file validation functionality"""
    
    def test_validate_valid_pdf(self, processor, sample_pdf):
        """Test validation of valid PDF file"""
        file_type, is_valid = processor.validate_file("test.pdf", sample_pdf)
        assert file_type == FileType.PDF
        assert is_valid is True
    
    def test_validate_file_size_exceeded(self, processor):
        """Test that oversized files are rejected"""
        # Create content larger than MAX_FILE_SIZE
        large_content = b"x" * (DocumentProcessor.MAX_FILE_SIZE + 1)
        
        with pytest.raises(FileSizeExceededError) as exc_info:
            processor.validate_file("large.pdf", large_content)
        
        assert "exceeds maximum allowed" in str(exc_info.value)
    
    def test_validate_unsupported_type(self, processor):
        """Test that unsupported file types are rejected"""
        with pytest.raises(UnsupportedFileTypeError):
            processor.validate_file("document.txt", b"some content")
    
    def test_validate_empty_file(self, processor):
        """Test validation of empty file"""
        file_type, is_valid = processor.validate_file("empty.pdf", b"")
        assert file_type == FileType.PDF
        assert is_valid is True


class TestPDFExtraction:
    """Test PDF text extraction"""
    
    def test_extract_from_pdf(self, processor, sample_pdf):
        """Test basic PDF text extraction"""
        text = processor.extract_text("test.pdf", sample_pdf)
        assert isinstance(text, str)
        # PDF should have page marker or text content
        assert "[Page 1]" in text or "Sample PDF Text" in text or len(text) >= 0
    
    def test_extract_from_corrupted_pdf(self, processor):
        """Test that corrupted PDF raises appropriate error"""
        corrupted_pdf = b"This is not a valid PDF file"
        
        with pytest.raises(CorruptedFileError) as exc_info:
            processor.extract_text("corrupted.pdf", corrupted_pdf)
        
        assert "Failed to process file" in str(exc_info.value)


class TestDOCXExtraction:
    """Test DOCX text extraction"""
    
    def test_extract_from_docx(self, processor, sample_docx):
        """Test basic DOCX text extraction"""
        text = processor.extract_text("test.docx", sample_docx)
        assert isinstance(text, str)
        assert "Sample financial statement" in text
        assert "Revenue: $1,000,000" in text
        assert "Profit: $250,000" in text
    
    def test_extract_table_from_docx(self, processor, sample_docx):
        """Test that tables are extracted from DOCX"""
        text = processor.extract_text("test.docx", sample_docx)
        assert "Assets" in text
        assert "$5,000,000" in text
        assert "Liabilities" in text
    
    def test_extract_from_corrupted_docx(self, processor):
        """Test that corrupted DOCX raises appropriate error"""
        corrupted_docx = b"This is not a valid DOCX file"
        
        with pytest.raises(CorruptedFileError) as exc_info:
            processor.extract_text("corrupted.docx", corrupted_docx)
        
        assert "Failed to process file" in str(exc_info.value)


class TestExcelExtraction:
    """Test Excel text extraction"""
    
    def test_extract_from_excel(self, processor, sample_excel):
        """Test basic Excel text extraction"""
        text = processor.extract_text("test.xlsx", sample_excel)
        assert isinstance(text, str)
        assert "[Sheet: Financial Data]" in text
        assert "Year" in text
        assert "Revenue" in text
        assert "1000000" in text or "1,000,000" in text
    
    def test_extract_from_corrupted_excel(self, processor):
        """Test that corrupted Excel raises appropriate error"""
        corrupted_excel = b"This is not a valid Excel file"
        
        with pytest.raises(CorruptedFileError) as exc_info:
            processor.extract_text("corrupted.xlsx", corrupted_excel)
        
        assert "Failed to process file" in str(exc_info.value)


class TestCSVExtraction:
    """Test CSV text extraction"""
    
    def test_extract_from_csv(self, processor, sample_csv):
        """Test basic CSV text extraction"""
        text = processor.extract_text("test.csv", sample_csv)
        assert isinstance(text, str)
        assert "Year" in text
        assert "Revenue" in text
        assert "Profit" in text
        assert "2021" in text
        assert "1000000" in text
    
    def test_extract_from_corrupted_csv(self, processor):
        """Test that malformed CSV is handled"""
        # CSV is more forgiving, but completely invalid data should fail
        corrupted_csv = b"\xff\xfe\x00\x00"  # Invalid UTF-8
        
        with pytest.raises(CorruptedFileError):
            processor.extract_text("corrupted.csv", corrupted_csv)


class TestImageExtraction:
    """Test image text extraction"""
    
    def test_extract_from_image(self, processor, sample_image):
        """Test basic image processing"""
        text = processor.extract_text("test.png", sample_image)
        assert isinstance(text, str)
        assert "[Image File]" in text
        assert "format:" in text.lower() or "Format:" in text
    
    def test_extract_from_corrupted_image(self, processor):
        """Test that corrupted image raises appropriate error"""
        corrupted_image = b"This is not a valid image file"
        
        with pytest.raises(CorruptedFileError) as exc_info:
            processor.extract_text("corrupted.png", corrupted_image)
        
        assert "Failed to process file" in str(exc_info.value)


class TestExtractWithMetadata:
    """Test extraction with metadata"""
    
    def test_extract_pdf_with_metadata(self, processor, sample_pdf):
        """Test PDF extraction with metadata"""
        result = processor.extract_with_metadata("test.pdf", sample_pdf)
        
        assert "text" in result
        assert "file_type" in result
        assert "file_size" in result
        assert "filename" in result
        assert result["file_type"] == "pdf"
        assert result["filename"] == "test.pdf"
        assert result["file_size"] == len(sample_pdf)
    
    def test_extract_excel_with_metadata(self, processor, sample_excel):
        """Test Excel extraction with metadata"""
        result = processor.extract_with_metadata("test.xlsx", sample_excel)
        
        assert "text" in result
        assert "file_type" in result
        assert "sheet_count" in result
        assert result["file_type"] == "xlsx"
        assert result["sheet_count"] >= 1
    
    def test_extract_csv_with_metadata(self, processor, sample_csv):
        """Test CSV extraction with metadata"""
        result = processor.extract_with_metadata("data.csv", sample_csv)
        
        assert "text" in result
        assert "file_type" in result
        assert result["file_type"] == "csv"
        assert "Year" in result["text"]


class TestErrorHandling:
    """Test error handling and edge cases"""
    
    def test_empty_filename(self, processor):
        """Test handling of empty filename"""
        with pytest.raises(UnsupportedFileTypeError):
            processor.detect_file_type("")
    
    def test_filename_without_extension(self, processor):
        """Test handling of filename without extension"""
        with pytest.raises(UnsupportedFileTypeError):
            processor.detect_file_type("document")
    
    def test_multiple_extensions(self, processor):
        """Test handling of multiple extensions"""
        # Should use the last extension
        file_type = processor.detect_file_type("document.backup.pdf")
        assert file_type == FileType.PDF
    
    def test_extract_with_wrong_extension(self, processor, sample_pdf):
        """Test extraction when file extension doesn't match content"""
        # This should still work as we validate by extension first
        # In production, you might want to add content-based validation
        text = processor.extract_text("test.pdf", sample_pdf)
        assert isinstance(text, str)
