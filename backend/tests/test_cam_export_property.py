"""
Property-based tests for CAM Export functionality

Feature: intelli-credit-platform
Property 16: CAM Export Format Validity

Validates: Requirements 7.4
"""

import pytest
from hypothesis import given, strategies as st, settings
from datetime import datetime
from unittest.mock import AsyncMock, MagicMock, patch
import io
from PyPDF2 import PdfReader
from docx import Document

from app.agents.cam_generator_agent import CAMGeneratorAgent
from app.models.domain import CAM


# Strategy for generating valid loan amounts
loan_amount_strategy = st.floats(
    min_value=100000.0,
    max_value=100000000.0,
    allow_nan=False,
    allow_infinity=False
)

# Strategy for generating credit scores (0-100)
credit_score_strategy = st.floats(
    min_value=0.0,
    max_value=100.0,
    allow_nan=False,
    allow_infinity=False
)

# Strategy for generating version numbers
version_strategy = st.integers(min_value=1, max_value=100)

# Strategy for generating company names
company_name_strategy = st.text(
    min_size=3,
    max_size=50,
    alphabet=st.characters(whitelist_categories=('Lu', 'Ll', 'Nd'), whitelist_characters=' &-.')
).filter(lambda x: x.strip() and not x.isspace())

# Strategy for generating loan purposes
loan_purpose_strategy = st.sampled_from([
    'Working capital',
    'Business expansion',
    'Equipment purchase',
    'Real estate acquisition',
    'Debt refinancing',
    'Inventory financing',
    'Technology upgrade',
    'Market expansion'
])

# Strategy for generating recommendations
recommendation_strategy = st.sampled_from(['approve', 'approve_with_conditions', 'reject'])

# Strategy for export formats
export_format_strategy = st.sampled_from(['pdf', 'word', 'docx'])


def create_analysis_data(
    application_id: str,
    company_name: str,
    loan_amount: float,
    loan_purpose: str,
    overall_score: float,
    recommendation: str,
    version: int
):
    """
    Helper function to create valid analysis data for CAM generation.
    
    This ensures we have properly structured data that the CAM generator expects.
    """
    return {
        'application_id': application_id,
        'company_name': company_name,
        'loan_amount': loan_amount,
        'loan_purpose': loan_purpose,
        'financial': {
            'ratios': {
                'current_ratio': {'value': 2.0, 'formatted_value': '2.00'},
                'debt_to_equity': {'value': 0.5, 'formatted_value': '0.50'},
                'roe': {'value': 0.15, 'formatted_value': '15.00%'},
                'net_profit_margin': {'value': 0.10, 'formatted_value': '10.00%'}
            },
            'benchmarks': {
                'current_ratio': {
                    'value': 2.0,
                    'benchmark_good': 1.5,
                    'performance': 'good'
                }
            },
            'trends': {
                'revenue': {
                    'trend_direction': 'increasing',
                    'cagr': 10.0,
                    'values': [1000000, 1100000, 1200000],
                    'interpretation': 'Positive growth trend'
                }
            },
            'summary': 'Financial analysis completed'
        },
        'forecasts': {
            'forecasts': {
                'revenue': {
                    'projected_values': [1300000, 1400000, 1500000],
                    'confidence_level': 0.75
                },
                'profit': {
                    'projected_values': [150000, 170000, 190000],
                    'confidence_level': 0.70
                }
            }
        },
        'risk': {
            'overall_score': overall_score,
            'recommendation': recommendation,
            'summary': f'Credit score: {overall_score:.1f}/100',
            'financial_health': {
                'score': overall_score,
                'weight': 0.35,
                'explanation': 'Financial health assessment',
                'key_findings': ['Strong ratios', 'Positive trends']
            },
            'cash_flow': {
                'score': overall_score,
                'weight': 0.25,
                'explanation': 'Cash flow assessment',
                'key_findings': ['Adequate liquidity']
            },
            'industry': {
                'score': overall_score,
                'weight': 0.15,
                'explanation': 'Industry assessment',
                'key_findings': ['Favorable outlook']
            },
            'promoter': {
                'score': overall_score,
                'weight': 0.15,
                'explanation': 'Promoter assessment',
                'key_findings': ['Experienced management']
            },
            'external_intelligence': {
                'score': overall_score,
                'weight': 0.10,
                'explanation': 'External intelligence assessment',
                'key_findings': ['Positive sentiment']
            }
        },
        'research': {
            'web': {
                'news_summary': 'Recent positive coverage',
                'market_sentiment': 'Positive'
            },
            'promoter': {
                'background': 'Experienced management team'
            },
            'industry': {
                'sector_outlook': 'Positive growth expected'
            }
        },
        'version': version
    }


class TestProperty16CAMExportFormatValidity:
    """
    Property 16: CAM Export Format Validity
    
    For any CAM export request specifying PDF or Word format, the system should
    generate a valid file in the requested format that can be opened by standard
    document readers.
    
    Validates: Requirements 7.4
    """
    
    @pytest.fixture(autouse=True)
    def setup_mock_openai(self):
        """Automatically mock OpenAI for all tests in this class."""
        with patch('app.agents.cam_generator_agent.AsyncOpenAI') as mock_client:
            # Create mock response object
            mock_response = MagicMock()
            mock_choice = MagicMock()
            mock_message = MagicMock()
            
            # Set up the response chain
            mock_message.content = "Generated content from OpenAI"
            mock_choice.message = mock_message
            mock_response.choices = [mock_choice]
            
            # Mock the async create method
            mock_create = AsyncMock(return_value=mock_response)
            mock_client.return_value.chat.completions.create = mock_create
            
            yield mock_client
    
    @pytest.mark.property
    @pytest.mark.asyncio
    @settings(max_examples=3)
    @given(
        company_name=company_name_strategy,
        loan_amount=loan_amount_strategy,
        loan_purpose=loan_purpose_strategy,
        overall_score=credit_score_strategy,
        recommendation=recommendation_strategy,
        version=version_strategy
    )
    async def test_pdf_export_produces_valid_pdf(
        self,
        company_name: str,
        loan_amount: float,
        loan_purpose: str,
        overall_score: float,
        recommendation: str,
        version: int
    ):
        """
        Property: PDF export must produce a valid PDF file
        
        For any CAM export request with format='pdf', the system should:
        - Return bytes data
        - Produce a valid PDF that can be parsed by PyPDF2
        - Contain at least one page
        - Contain extractable text content
        """
        agent = CAMGeneratorAgent()
        
        # Create analysis data and generate CAM
        analysis_data = create_analysis_data(
            application_id=f'test-app-{version}',
            company_name=company_name,
            loan_amount=loan_amount,
            loan_purpose=loan_purpose,
            overall_score=overall_score,
            recommendation=recommendation,
            version=version
        )
        
        cam = await agent.generate(analysis_data)
        
        # Export to PDF
        pdf_bytes = await agent.export_to_pdf(cam)
        
        # Verify output is bytes
        assert isinstance(pdf_bytes, bytes), \
            f"PDF export should return bytes, got {type(pdf_bytes)}"
        
        # Verify bytes are non-empty
        assert len(pdf_bytes) > 0, \
            "PDF export produced empty bytes"
        
        # Verify PDF is valid by parsing it
        try:
            pdf_stream = io.BytesIO(pdf_bytes)
            pdf_reader = PdfReader(pdf_stream)
            
            # Verify PDF has at least one page
            assert len(pdf_reader.pages) > 0, \
                "PDF should have at least one page"
            
            # Verify PDF contains text (extract from first page)
            first_page = pdf_reader.pages[0]
            text = first_page.extract_text()
            
            assert text is not None, \
                "PDF should contain extractable text"
            assert len(text.strip()) > 0, \
                "PDF text content should not be empty"
            
            # Verify PDF contains key information from CAM
            # (company name should appear in the PDF)
            pdf_full_text = ""
            for page in pdf_reader.pages:
                pdf_full_text += page.extract_text()
            
            assert company_name in pdf_full_text, \
                f"PDF should contain company name '{company_name}'"
            
        except Exception as e:
            pytest.fail(f"Failed to parse PDF: {str(e)}")
    
    @pytest.mark.property
    @pytest.mark.asyncio
    @settings(max_examples=3)
    @given(
        company_name=company_name_strategy,
        loan_amount=loan_amount_strategy,
        loan_purpose=loan_purpose_strategy,
        overall_score=credit_score_strategy,
        recommendation=recommendation_strategy,
        version=version_strategy
    )
    async def test_word_export_produces_valid_docx(
        self,
        company_name: str,
        loan_amount: float,
        loan_purpose: str,
        overall_score: float,
        recommendation: str,
        version: int
    ):
        """
        Property: Word export must produce a valid DOCX file
        
        For any CAM export request with format='word' or 'docx', the system should:
        - Return bytes data
        - Produce a valid DOCX that can be opened by python-docx
        - Contain paragraphs with text content
        - Contain key information from the CAM
        """
        agent = CAMGeneratorAgent()
        
        # Create analysis data and generate CAM
        analysis_data = create_analysis_data(
            application_id=f'test-app-{version}',
            company_name=company_name,
            loan_amount=loan_amount,
            loan_purpose=loan_purpose,
            overall_score=overall_score,
            recommendation=recommendation,
            version=version
        )
        
        cam = await agent.generate(analysis_data)
        
        # Export to Word
        docx_bytes = await agent.export_to_word(cam)
        
        # Verify output is bytes
        assert isinstance(docx_bytes, bytes), \
            f"Word export should return bytes, got {type(docx_bytes)}"
        
        # Verify bytes are non-empty
        assert len(docx_bytes) > 0, \
            "Word export produced empty bytes"
        
        # Verify DOCX is valid by opening it
        try:
            docx_stream = io.BytesIO(docx_bytes)
            doc = Document(docx_stream)
            
            # Verify document has paragraphs
            assert len(doc.paragraphs) > 0, \
                "DOCX should have at least one paragraph"
            
            # Extract all text from document
            doc_text = ""
            for paragraph in doc.paragraphs:
                doc_text += paragraph.text + "\n"
            
            # Verify document contains text
            assert len(doc_text.strip()) > 0, \
                "DOCX should contain text content"
            
            # Verify document contains key information from CAM
            assert company_name in doc_text, \
                f"DOCX should contain company name '{company_name}'"
            
            # Verify document has core properties set
            assert doc.core_properties.title is not None, \
                "DOCX should have title property set"
            assert doc.core_properties.created is not None, \
                "DOCX should have created date property set"
            
        except Exception as e:
            pytest.fail(f"Failed to parse DOCX: {str(e)}")
    
    @pytest.mark.property
    @pytest.mark.asyncio
    @settings(max_examples=3)
    @given(
        company_name=company_name_strategy,
        loan_amount=loan_amount_strategy,
        loan_purpose=loan_purpose_strategy,
        overall_score=credit_score_strategy,
        recommendation=recommendation_strategy,
        version=version_strategy,
        export_format=export_format_strategy
    )
    async def test_export_format_validation(
        self,
        company_name: str,
        loan_amount: float,
        loan_purpose: str,
        overall_score: float,
        recommendation: str,
        version: int,
        export_format: str
    ):
        """
        Property: Export format validation must accept valid formats
        
        For any valid export format ('pdf', 'word', 'docx'), the validation
        should return True. The validation is case-insensitive.
        """
        agent = CAMGeneratorAgent()
        
        # Test format validation
        is_valid = agent.validate_export_format(export_format)
        
        # All formats in our strategy should be valid
        assert is_valid is True, \
            f"Format '{export_format}' should be valid"
        
        # Test case-insensitivity
        is_valid_upper = agent.validate_export_format(export_format.upper())
        assert is_valid_upper is True, \
            f"Format '{export_format.upper()}' should be valid (case-insensitive)"
        
        is_valid_mixed = agent.validate_export_format(export_format.title())
        assert is_valid_mixed is True, \
            f"Format '{export_format.title()}' should be valid (case-insensitive)"
    
    @pytest.mark.property
    @pytest.mark.asyncio
    @settings(max_examples=3)
    @given(
        invalid_format=st.text(
            min_size=1,
            max_size=20,
            alphabet=st.characters(whitelist_categories=('Lu', 'Ll', 'Nd'))
        ).filter(lambda x: x.lower() not in ['pdf', 'word', 'docx'])
    )
    async def test_export_format_validation_rejects_invalid(
        self,
        invalid_format: str
    ):
        """
        Property: Export format validation must reject invalid formats
        
        For any format that is not 'pdf', 'word', or 'docx', the validation
        should return False.
        """
        agent = CAMGeneratorAgent()
        
        # Test format validation with invalid format
        is_valid = agent.validate_export_format(invalid_format)
        
        # Invalid formats should be rejected
        assert is_valid is False, \
            f"Format '{invalid_format}' should be invalid"
    
    @pytest.mark.property
    @pytest.mark.asyncio
    @settings(max_examples=3)
    @given(
        company_name=company_name_strategy,
        loan_amount=loan_amount_strategy,
        loan_purpose=loan_purpose_strategy,
        overall_score=credit_score_strategy,
        recommendation=recommendation_strategy,
        version=version_strategy
    )
    async def test_pdf_export_contains_cam_metadata(
        self,
        company_name: str,
        loan_amount: float,
        loan_purpose: str,
        overall_score: float,
        recommendation: str,
        version: int
    ):
        """
        Property: PDF export must contain CAM metadata
        
        For any PDF export, the document should contain:
        - Company name
        - Loan amount
        - Credit score
        - Version number
        - Generation timestamp
        """
        agent = CAMGeneratorAgent()
        
        # Create analysis data and generate CAM
        analysis_data = create_analysis_data(
            application_id=f'test-app-{version}',
            company_name=company_name,
            loan_amount=loan_amount,
            loan_purpose=loan_purpose,
            overall_score=overall_score,
            recommendation=recommendation,
            version=version
        )
        
        cam = await agent.generate(analysis_data)
        
        # Export to PDF
        pdf_bytes = await agent.export_to_pdf(cam)
        
        # Parse PDF and extract all text
        pdf_stream = io.BytesIO(pdf_bytes)
        pdf_reader = PdfReader(pdf_stream)
        
        pdf_full_text = ""
        for page in pdf_reader.pages:
            pdf_full_text += page.extract_text()
        
        # Verify company name is present
        assert company_name in pdf_full_text, \
            f"PDF should contain company name '{company_name}'"
        
        # Verify loan amount is present (in some format)
        loan_amount_int = int(loan_amount)
        assert (
            str(loan_amount_int) in pdf_full_text or
            f"{loan_amount:,.2f}" in pdf_full_text or
            f"{loan_amount_int:,}" in pdf_full_text
        ), f"PDF should contain loan amount {loan_amount}"
        
        # Verify credit score is present
        score_str = f"{overall_score:.1f}"
        score_int = f"{int(overall_score)}"
        assert score_str in pdf_full_text or score_int in pdf_full_text, \
            f"PDF should contain credit score {overall_score}"
        
        # Verify version is present
        assert str(version) in pdf_full_text, \
            f"PDF should contain version number {version}"
    
    @pytest.mark.property
    @pytest.mark.asyncio
    @settings(max_examples=3)
    @given(
        company_name=company_name_strategy,
        loan_amount=loan_amount_strategy,
        loan_purpose=loan_purpose_strategy,
        overall_score=credit_score_strategy,
        recommendation=recommendation_strategy,
        version=version_strategy
    )
    async def test_word_export_contains_cam_metadata(
        self,
        company_name: str,
        loan_amount: float,
        loan_purpose: str,
        overall_score: float,
        recommendation: str,
        version: int
    ):
        """
        Property: Word export must contain CAM metadata
        
        For any Word export, the document should contain:
        - Company name
        - Loan amount
        - Credit score
        - Version number
        - Document properties (title, created date)
        """
        agent = CAMGeneratorAgent()
        
        # Create analysis data and generate CAM
        analysis_data = create_analysis_data(
            application_id=f'test-app-{version}',
            company_name=company_name,
            loan_amount=loan_amount,
            loan_purpose=loan_purpose,
            overall_score=overall_score,
            recommendation=recommendation,
            version=version
        )
        
        cam = await agent.generate(analysis_data)
        
        # Export to Word
        docx_bytes = await agent.export_to_word(cam)
        
        # Parse DOCX and extract all text
        docx_stream = io.BytesIO(docx_bytes)
        doc = Document(docx_stream)
        
        doc_text = ""
        for paragraph in doc.paragraphs:
            doc_text += paragraph.text + "\n"
        
        # Verify company name is present
        assert company_name in doc_text, \
            f"DOCX should contain company name '{company_name}'"
        
        # Verify loan amount is present (in some format)
        loan_amount_int = int(loan_amount)
        assert (
            str(loan_amount_int) in doc_text or
            f"{loan_amount:,.2f}" in doc_text or
            f"{loan_amount_int:,}" in doc_text
        ), f"DOCX should contain loan amount {loan_amount}"
        
        # Verify credit score is present
        score_str = f"{overall_score:.1f}"
        score_int = f"{int(overall_score)}"
        assert score_str in doc_text or score_int in doc_text, \
            f"DOCX should contain credit score {overall_score}"
        
        # Verify version is present
        assert str(version) in doc_text, \
            f"DOCX should contain version number {version}"
        
        # Verify document properties
        assert doc.core_properties.title == "Credit Appraisal Memo", \
            "DOCX should have correct title property"
        
        assert doc.core_properties.created == cam.generated_at, \
            "DOCX created date should match CAM generation timestamp"
    
    @pytest.mark.property
    @pytest.mark.asyncio
    @settings(max_examples=3)
    @given(
        company_name=company_name_strategy,
        loan_amount=loan_amount_strategy,
        loan_purpose=loan_purpose_strategy,
        overall_score=credit_score_strategy,
        recommendation=recommendation_strategy,
        version=version_strategy
    )
    async def test_pdf_export_is_deterministic_for_same_cam(
        self,
        company_name: str,
        loan_amount: float,
        loan_purpose: str,
        overall_score: float,
        recommendation: str,
        version: int
    ):
        """
        Property: PDF export should be consistent for the same CAM
        
        For any CAM, exporting to PDF multiple times should produce
        files with the same content (though byte-level differences may exist
        due to timestamps or metadata).
        """
        agent = CAMGeneratorAgent()
        
        # Create analysis data and generate CAM
        analysis_data = create_analysis_data(
            application_id=f'test-app-{version}',
            company_name=company_name,
            loan_amount=loan_amount,
            loan_purpose=loan_purpose,
            overall_score=overall_score,
            recommendation=recommendation,
            version=version
        )
        
        cam = await agent.generate(analysis_data)
        
        # Export to PDF twice
        pdf_bytes_1 = await agent.export_to_pdf(cam)
        pdf_bytes_2 = await agent.export_to_pdf(cam)
        
        # Parse both PDFs and extract text
        pdf_stream_1 = io.BytesIO(pdf_bytes_1)
        pdf_reader_1 = PdfReader(pdf_stream_1)
        text_1 = ""
        for page in pdf_reader_1.pages:
            text_1 += page.extract_text()
        
        pdf_stream_2 = io.BytesIO(pdf_bytes_2)
        pdf_reader_2 = PdfReader(pdf_stream_2)
        text_2 = ""
        for page in pdf_reader_2.pages:
            text_2 += page.extract_text()
        
        # Verify both PDFs have the same number of pages
        assert len(pdf_reader_1.pages) == len(pdf_reader_2.pages), \
            "Both PDF exports should have the same number of pages"
        
        # Verify both PDFs have the same text content
        assert text_1 == text_2, \
            "Both PDF exports should have the same text content"
    
    @pytest.mark.property
    @pytest.mark.asyncio
    @settings(max_examples=3)
    @given(
        company_name=company_name_strategy,
        loan_amount=loan_amount_strategy,
        loan_purpose=loan_purpose_strategy,
        overall_score=credit_score_strategy,
        recommendation=recommendation_strategy,
        version=version_strategy
    )
    async def test_word_export_is_deterministic_for_same_cam(
        self,
        company_name: str,
        loan_amount: float,
        loan_purpose: str,
        overall_score: float,
        recommendation: str,
        version: int
    ):
        """
        Property: Word export should be consistent for the same CAM
        
        For any CAM, exporting to Word multiple times should produce
        documents with the same content structure.
        """
        agent = CAMGeneratorAgent()
        
        # Create analysis data and generate CAM
        analysis_data = create_analysis_data(
            application_id=f'test-app-{version}',
            company_name=company_name,
            loan_amount=loan_amount,
            loan_purpose=loan_purpose,
            overall_score=overall_score,
            recommendation=recommendation,
            version=version
        )
        
        cam = await agent.generate(analysis_data)
        
        # Export to Word twice
        docx_bytes_1 = await agent.export_to_word(cam)
        docx_bytes_2 = await agent.export_to_word(cam)
        
        # Parse both DOCX files and extract text
        docx_stream_1 = io.BytesIO(docx_bytes_1)
        doc_1 = Document(docx_stream_1)
        text_1 = ""
        for paragraph in doc_1.paragraphs:
            text_1 += paragraph.text + "\n"
        
        docx_stream_2 = io.BytesIO(docx_bytes_2)
        doc_2 = Document(docx_stream_2)
        text_2 = ""
        for paragraph in doc_2.paragraphs:
            text_2 += paragraph.text + "\n"
        
        # Verify both documents have the same number of paragraphs
        assert len(doc_1.paragraphs) == len(doc_2.paragraphs), \
            "Both Word exports should have the same number of paragraphs"
        
        # Verify both documents have the same text content
        assert text_1 == text_2, \
            "Both Word exports should have the same text content"
        
        # Verify both documents have the same properties
        assert doc_1.core_properties.title == doc_2.core_properties.title, \
            "Both Word exports should have the same title"
        assert doc_1.core_properties.created == doc_2.core_properties.created, \
            "Both Word exports should have the same created date"
