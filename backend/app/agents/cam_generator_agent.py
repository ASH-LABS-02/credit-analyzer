"""
CAM Generator Agent

This agent compiles all analysis results into a structured Credit Appraisal Memo (CAM)
with professional formatting, tables, charts, and comprehensive documentation.

Requirements: 7.1, 7.2, 7.3, 7.4, 7.5
"""

import json
import io
from datetime import datetime
from typing import Dict, Any, List, Literal, Optional
from openai import AsyncOpenAI
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2
from weasyprint import HTML, CSS

from app.core.config import settings
from app.core.audit_logger import AuditLogger
from app.models.domain import CAM, RiskAssessment, Forecast, ResearchFindings


class CAMGeneratorAgent:
    """
    AI agent for generating Credit Appraisal Memos (CAMs).
    
    Compiles all analysis results into a structured, professional CAM document with:
    - Executive summary
    - Company overview
    - Financial analysis with tables and charts
    - Risk assessment with detailed scoring
    - Credit recommendation with supporting rationale
    - Professional formatting and version tracking
    
    Requirements: 7.1, 7.2, 7.3, 7.5
    """
    
    def __init__(self, audit_logger: Optional[AuditLogger] = None):
        """
        Initialize the CAM Generator Agent.
        
        Args:
            audit_logger: Optional audit logger for AI decision logging
        """
        self.openai = AsyncOpenAI(api_key=settings.OPENAI_API_KEY)
        self.audit_logger = audit_logger
    
    async def generate(self, analysis_data: Dict[str, Any]) -> CAM:
        """
        Generate a comprehensive Credit Appraisal Memo from analysis results.
        
        Args:
            analysis_data: Dictionary containing:
                - application_id: Application identifier
                - company_name: Company name
                - loan_amount: Requested loan amount
                - loan_purpose: Purpose of the loan
                - financial: Financial analysis results from FinancialAnalysisAgent
                - forecasts: Forecast results from ForecastingAgent
                - risk: Risk assessment from RiskScoringAgent
                - research: Research findings from research agents
                - version: Optional version number (defaults to 1)
        
        Returns:
            CAM object with structured content and metadata
        
        Requirements: 7.1, 7.2, 7.3, 7.5
        """
        # Extract data components
        application_id = analysis_data.get('application_id', '')
        company_name = analysis_data.get('company_name', 'Unknown Company')
        loan_amount = analysis_data.get('loan_amount', 0)
        loan_purpose = analysis_data.get('loan_purpose', 'Not specified')
        financial_data = analysis_data.get('financial', {})
        forecasts_data = analysis_data.get('forecasts', {})
        risk_data = analysis_data.get('risk', {})
        research_data = analysis_data.get('research', {})
        version = analysis_data.get('version', 1)
        
        # Generate each section
        executive_summary = await self._generate_executive_summary(
            company_name, loan_amount, risk_data, financial_data
        )
        
        company_overview = await self._generate_company_overview(
            company_name, loan_purpose, research_data
        )
        
        financial_analysis = await self._generate_financial_analysis(
            financial_data, forecasts_data
        )
        
        risk_assessment = await self._generate_risk_assessment(
            risk_data
        )
        
        recommendation = await self._generate_recommendation(
            risk_data, loan_amount, financial_data
        )
        
        # Build sections dictionary
        sections = {
            'executive_summary': executive_summary,
            'company_overview': company_overview,
            'financial_analysis': financial_analysis,
            'risk_assessment': risk_assessment,
            'recommendation': recommendation
        }
        
        # Compile full CAM content
        content = self._compile_cam_content(
            company_name, loan_amount, sections, version
        )
        
        # Create CAM object
        cam = CAM(
            application_id=application_id,
            content=content,
            version=version,
            generated_at=datetime.utcnow(),
            sections=sections
        )
        
        # Log AI decision
        if self.audit_logger:
            try:
                await self.audit_logger.log_ai_decision(
                    agent_name='CAMGeneratorAgent',
                    application_id=application_id,
                    decision=f"Generated Credit Appraisal Memo v{version} for {company_name}",
                    reasoning=f"Compiled comprehensive CAM report with {len(sections)} sections: "
                             f"{', '.join(sections.keys())}. "
                             f"Loan amount: ${loan_amount:,.2f}. "
                             f"Credit recommendation: {risk_data.get('recommendation', 'unknown')}.",
                    data_sources=['financial_analysis', 'forecasts', 'risk_assessment', 'research_findings'],
                    additional_details={
                        'version': version,
                        'sections_count': len(sections),
                        'sections': list(sections.keys()),
                        'company_name': company_name,
                        'loan_amount': loan_amount,
                        'content_length': len(content)
                    }
                )
            except Exception as e:
                print(f"Error logging AI decision: {str(e)}")
        
        return cam

    async def _generate_executive_summary(
        self,
        company_name: str,
        loan_amount: float,
        risk_data: Dict[str, Any],
        financial_data: Dict[str, Any]
    ) -> str:
        """
        Generate executive summary section.
        
        Args:
            company_name: Company name
            loan_amount: Requested loan amount
            risk_data: Risk assessment data
            financial_data: Financial analysis data
        
        Returns:
            Executive summary content in Markdown format
        
        Requirements: 7.2
        """
        # Extract key information
        overall_score = risk_data.get('overall_score', 0)
        recommendation = risk_data.get('recommendation', 'unknown')
        summary_text = risk_data.get('summary', '')
        
        # Build prompt for AI generation
        prompt = f"""Generate a concise executive summary for a Credit Appraisal Memo.

**Company:** {company_name}
**Loan Amount:** ${loan_amount:,.2f}
**Credit Score:** {overall_score:.1f}/100
**Recommendation:** {recommendation}

**Risk Assessment Summary:**
{summary_text}

**Financial Highlights:**
{json.dumps(financial_data.get('summary', 'Financial analysis completed'), indent=2)}

Create a professional executive summary (3-4 paragraphs) that:
1. States the loan request and credit recommendation upfront
2. Summarizes the company's financial position
3. Highlights key strengths and concerns
4. Provides the overall credit assessment

Use professional, clear language suitable for senior management review.
Return ONLY the summary text in Markdown format, no additional formatting.
"""
        
        try:
            response = await self.openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a senior credit analyst writing executive summaries for Credit Appraisal Memos. Be concise, professional, and decision-focused."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=500
            )
            
            return response.choices[0].message.content.strip()
        
        except Exception as e:
            print(f"Error generating executive summary: {str(e)}")
            return self._generate_fallback_executive_summary(
                company_name, loan_amount, overall_score, recommendation
            )

    async def _generate_company_overview(
        self,
        company_name: str,
        loan_purpose: str,
        research_data: Dict[str, Any]
    ) -> str:
        """
        Generate company overview section.
        
        Args:
            company_name: Company name
            loan_purpose: Purpose of the loan
            research_data: Research findings
        
        Returns:
            Company overview content in Markdown format
        
        Requirements: 7.2
        """
        # Extract research components
        web_research = research_data.get('web', {}).get('news_summary', 'No web research available')
        promoter_analysis = research_data.get('promoter', {}).get('background', 'No promoter analysis available')
        industry_analysis = research_data.get('industry', {}).get('sector_outlook', 'No industry analysis available')
        
        prompt = f"""Generate a comprehensive company overview section for a Credit Appraisal Memo.

**Company:** {company_name}
**Loan Purpose:** {loan_purpose}

**Web Research:**
{web_research}

**Promoter/Management Analysis:**
{promoter_analysis}

**Industry Context:**
{industry_analysis}

Create a professional company overview (4-5 paragraphs) covering:
1. Company background and business model
2. Management team and promoter profile
3. Industry position and competitive landscape
4. Recent developments and market presence
5. Purpose and use of requested loan

Use clear, professional language with proper structure.
Return ONLY the overview text in Markdown format with appropriate headings.
"""
        
        try:
            response = await self.openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a credit analyst writing company overview sections. Be thorough, objective, and well-structured."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=800
            )
            
            return response.choices[0].message.content.strip()
        
        except Exception as e:
            print(f"Error generating company overview: {str(e)}")
            return self._generate_fallback_company_overview(
                company_name, loan_purpose, research_data
            )

    async def _generate_financial_analysis(
        self,
        financial_data: Dict[str, Any],
        forecasts_data: Dict[str, Any]
    ) -> str:
        """
        Generate financial analysis section with tables and charts.
        
        Args:
            financial_data: Financial analysis results
            forecasts_data: Forecast results
        
        Returns:
            Financial analysis content in Markdown format with tables
        
        Requirements: 7.2, 7.3
        """
        # Extract financial components
        ratios = financial_data.get('ratios', {})
        trends = financial_data.get('trends', {})
        benchmarks = financial_data.get('benchmarks', {})
        forecasts = forecasts_data.get('forecasts', {})
        
        # Build financial ratios table
        ratios_table = self._build_ratios_table(ratios, benchmarks)
        
        # Build trends table
        trends_table = self._build_trends_table(trends)
        
        # Build forecasts table
        forecasts_table = self._build_forecasts_table(forecasts)
        
        # Generate narrative analysis
        prompt = f"""Generate a comprehensive financial analysis narrative for a Credit Appraisal Memo.

**Financial Ratios:**
{json.dumps(ratios, indent=2, default=str)}

**Trends:**
{json.dumps(trends, indent=2, default=str)}

**Benchmarks:**
{json.dumps(benchmarks, indent=2, default=str)}

**Forecasts:**
{json.dumps(forecasts, indent=2, default=str)}

Create a professional financial analysis (5-6 paragraphs) covering:
1. Overview of financial position and health
2. Liquidity and working capital analysis
3. Leverage and capital structure assessment
4. Profitability and efficiency analysis
5. Historical trends and growth patterns
6. Future outlook based on forecasts

Use clear, analytical language with specific metrics and comparisons.
Return ONLY the narrative text in Markdown format.
"""
        
        try:
            response = await self.openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a financial analyst writing detailed financial analysis sections. Be thorough, data-driven, and analytical."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=1000
            )
            
            narrative = response.choices[0].message.content.strip()
        
        except Exception as e:
            print(f"Error generating financial analysis narrative: {str(e)}")
            narrative = self._generate_fallback_financial_narrative(
                ratios, trends, benchmarks
            )
        
        # Combine narrative with tables
        content = f"""## Financial Analysis

{narrative}

### Key Financial Ratios

{ratios_table}

### Historical Trends

{trends_table}

### Financial Forecasts (3-Year Projection)

{forecasts_table}
"""
        
        return content

    async def _generate_risk_assessment(
        self,
        risk_data: Dict[str, Any]
    ) -> str:
        """
        Generate risk assessment section with factor breakdown.
        
        Args:
            risk_data: Risk assessment data
        
        Returns:
            Risk assessment content in Markdown format
        
        Requirements: 7.2, 7.3
        """
        # Extract risk components
        overall_score = risk_data.get('overall_score', 0)
        recommendation = risk_data.get('recommendation', 'unknown')
        
        # Extract individual factor scores
        financial_health = risk_data.get('financial_health', {})
        cash_flow = risk_data.get('cash_flow', {})
        industry = risk_data.get('industry', {})
        promoter = risk_data.get('promoter', {})
        external_intelligence = risk_data.get('external_intelligence', {})
        
        # Build risk factors table
        risk_factors_table = self._build_risk_factors_table(
            financial_health, cash_flow, industry, promoter, external_intelligence
        )
        
        # Generate narrative
        prompt = f"""Generate a comprehensive risk assessment narrative for a Credit Appraisal Memo.

**Overall Credit Score:** {overall_score:.1f}/100
**Recommendation:** {recommendation}

**Risk Factor Scores:**
- Financial Health: {financial_health.get('score', 0):.1f}/100 (Weight: 35%)
- Cash Flow: {cash_flow.get('score', 0):.1f}/100 (Weight: 25%)
- Industry: {industry.get('score', 0):.1f}/100 (Weight: 15%)
- Promoter: {promoter.get('score', 0):.1f}/100 (Weight: 15%)
- External Intelligence: {external_intelligence.get('score', 0):.1f}/100 (Weight: 10%)

**Factor Explanations:**
Financial Health: {financial_health.get('explanation', 'N/A')}
Cash Flow: {cash_flow.get('explanation', 'N/A')}
Industry: {industry.get('explanation', 'N/A')}
Promoter: {promoter.get('explanation', 'N/A')}
External Intelligence: {external_intelligence.get('explanation', 'N/A')}

Create a professional risk assessment (4-5 paragraphs) covering:
1. Overall credit risk profile and score interpretation
2. Analysis of each risk factor and its contribution
3. Key risk drivers and mitigating factors
4. Comparison of strengths vs. weaknesses
5. Overall risk conclusion

Use clear, analytical language focused on credit risk.
Return ONLY the narrative text in Markdown format.
"""
        
        try:
            response = await self.openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a credit risk analyst writing risk assessment sections. Be thorough, balanced, and risk-focused."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=800
            )
            
            narrative = response.choices[0].message.content.strip()
        
        except Exception as e:
            print(f"Error generating risk assessment narrative: {str(e)}")
            narrative = self._generate_fallback_risk_narrative(
                overall_score, recommendation, financial_health, cash_flow
            )
        
        # Combine narrative with table
        content = f"""## Risk Assessment

{narrative}

### Risk Factor Breakdown

{risk_factors_table}

**Overall Credit Score:** {overall_score:.1f}/100

**Credit Recommendation:** {recommendation.upper().replace('_', ' ')}
"""
        
        return content

    async def _generate_recommendation(
        self,
        risk_data: Dict[str, Any],
        loan_amount: float,
        financial_data: Dict[str, Any]
    ) -> str:
        """
        Generate credit recommendation section.
        
        Args:
            risk_data: Risk assessment data
            loan_amount: Requested loan amount
            financial_data: Financial analysis data
        
        Returns:
            Recommendation content in Markdown format
        
        Requirements: 7.2
        """
        overall_score = risk_data.get('overall_score', 0)
        recommendation = risk_data.get('recommendation', 'unknown')
        summary = risk_data.get('summary', '')
        
        # Extract key findings from risk factors
        key_findings = []
        for factor_name in ['financial_health', 'cash_flow', 'industry', 'promoter', 'external_intelligence']:
            factor = risk_data.get(factor_name, {})
            findings = factor.get('key_findings', [])
            key_findings.extend(findings[:2])  # Top 2 from each factor
        
        prompt = f"""Generate a comprehensive credit recommendation section for a Credit Appraisal Memo.

**Loan Amount:** ${loan_amount:,.2f}
**Credit Score:** {overall_score:.1f}/100
**Recommendation:** {recommendation}

**Risk Summary:**
{summary}

**Key Findings:**
{chr(10).join(f'- {finding}' for finding in key_findings[:10])}

Create a professional recommendation section (3-4 paragraphs) that:
1. States the clear credit recommendation (Approve/Approve with Conditions/Reject)
2. Provides the rationale based on the analysis
3. If conditional approval, specifies the conditions and covenants
4. If rejection, explains the key concerns
5. Suggests monitoring requirements or next steps

Use decisive, clear language appropriate for credit committee review.
Return ONLY the recommendation text in Markdown format.
"""
        
        try:
            response = await self.openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a senior credit officer writing credit recommendations. Be clear, decisive, and actionable."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=600
            )
            
            return response.choices[0].message.content.strip()
        
        except Exception as e:
            print(f"Error generating recommendation: {str(e)}")
            return self._generate_fallback_recommendation(
                recommendation, overall_score, loan_amount
            )

    def _compile_cam_content(
        self,
        company_name: str,
        loan_amount: float,
        sections: Dict[str, str],
        version: int
    ) -> str:
        """
        Compile all sections into a complete CAM document.
        
        Args:
            company_name: Company name
            loan_amount: Requested loan amount
            sections: Dictionary of section content
            version: CAM version number
        
        Returns:
            Complete CAM content in Markdown format
        
        Requirements: 7.3, 7.5
        """
        timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
        
        content = f"""# Credit Appraisal Memo

**Company:** {company_name}  
**Loan Amount:** ${loan_amount:,.2f}  
**Generated:** {timestamp}  
**Version:** {version}

---

## Executive Summary

{sections['executive_summary']}

---

## Company Overview

{sections['company_overview']}

---

{sections['financial_analysis']}

---

{sections['risk_assessment']}

---

## Credit Recommendation

{sections['recommendation']}

---

*This Credit Appraisal Memo was generated by the Intelli-Credit AI-Powered Credit Decisioning Platform.*  
*Document Version: {version} | Generated: {timestamp}*
"""
        
        return content

    def _build_ratios_table(
        self,
        ratios: Dict[str, Any],
        benchmarks: Dict[str, Any]
    ) -> str:
        """
        Build a formatted table of financial ratios.
        
        Args:
            ratios: Calculated ratios
            benchmarks: Benchmark comparisons
        
        Returns:
            Markdown table of ratios
        
        Requirements: 7.3
        """
        if not ratios:
            return "*No financial ratios available.*"
        
        table = "| Ratio | Value | Industry Benchmark | Performance |\n"
        table += "|-------|-------|-------------------|-------------|\n"
        
        for ratio_name, ratio_data in ratios.items():
            value = ratio_data.get('formatted_value', 'N/A')
            
            # Get benchmark data if available
            benchmark_data = benchmarks.get(ratio_name, {})
            if benchmark_data:
                benchmark_good = benchmark_data.get('benchmark_good', 'N/A')
                performance = benchmark_data.get('performance', 'N/A')
                
                # Format benchmark display
                if isinstance(benchmark_good, float):
                    if ratio_name in ['net_profit_margin', 'roe', 'roa']:
                        benchmark_str = f"{benchmark_good * 100:.1f}%"
                    else:
                        benchmark_str = f"{benchmark_good:.2f}"
                else:
                    benchmark_str = str(benchmark_good)
                
                # Add emoji for performance
                if performance == 'good':
                    performance_str = "✓ Good"
                elif performance == 'acceptable':
                    performance_str = "~ Acceptable"
                else:
                    performance_str = "✗ Below"
            else:
                benchmark_str = "N/A"
                performance_str = "N/A"
            
            # Format ratio name
            ratio_display = ratio_name.replace('_', ' ').title()
            
            table += f"| {ratio_display} | {value} | {benchmark_str} | {performance_str} |\n"
        
        return table

    def _build_trends_table(self, trends: Dict[str, Any]) -> str:
        """
        Build a formatted table of historical trends.
        
        Args:
            trends: Trend analysis data
        
        Returns:
            Markdown table of trends
        
        Requirements: 7.3
        """
        if not trends:
            return "*No trend data available.*"
        
        table = "| Metric | Trend Direction | CAGR | Interpretation |\n"
        table += "|--------|----------------|------|----------------|\n"
        
        for metric_name, trend_data in trends.items():
            metric_display = metric_name.replace('_', ' ').title()
            trend_direction = trend_data.get('trend_direction', 'unknown')
            cagr = trend_data.get('cagr')
            
            # Format CAGR
            if cagr is not None:
                cagr_str = f"{cagr:.1f}%"
            else:
                cagr_str = "N/A"
            
            # Add emoji for trend direction
            if trend_direction == 'increasing':
                trend_str = "↑ Increasing"
            elif trend_direction == 'decreasing':
                trend_str = "↓ Decreasing"
            elif trend_direction == 'stable':
                trend_str = "→ Stable"
            else:
                trend_str = "~ Volatile"
            
            # Get interpretation (truncate if too long)
            interpretation = trend_data.get('interpretation', 'N/A')
            if len(interpretation) > 60:
                interpretation = interpretation[:57] + "..."
            
            table += f"| {metric_display} | {trend_str} | {cagr_str} | {interpretation} |\n"
        
        return table

    def _build_forecasts_table(self, forecasts: Dict[str, Any]) -> str:
        """
        Build a formatted table of financial forecasts.
        
        Args:
            forecasts: Forecast data
        
        Returns:
            Markdown table of forecasts
        
        Requirements: 7.3
        """
        if not forecasts:
            return "*No forecast data available.*"
        
        table = "| Metric | Year 1 | Year 2 | Year 3 | Confidence |\n"
        table += "|--------|--------|--------|--------|------------|\n"
        
        for metric_name, forecast_data in forecasts.items():
            metric_display = metric_name.replace('_', ' ').title()
            projected_values = forecast_data.get('projected_values', [])
            confidence = forecast_data.get('confidence_level', 0)
            
            # Format projected values
            if len(projected_values) >= 3:
                year1 = f"${projected_values[0]:,.0f}" if projected_values[0] >= 1000 else f"{projected_values[0]:.2f}"
                year2 = f"${projected_values[1]:,.0f}" if projected_values[1] >= 1000 else f"{projected_values[1]:.2f}"
                year3 = f"${projected_values[2]:,.0f}" if projected_values[2] >= 1000 else f"{projected_values[2]:.2f}"
            else:
                year1 = year2 = year3 = "N/A"
            
            # Format confidence
            confidence_str = f"{confidence * 100:.0f}%" if confidence else "N/A"
            
            table += f"| {metric_display} | {year1} | {year2} | {year3} | {confidence_str} |\n"
        
        return table

    def _build_risk_factors_table(
        self,
        financial_health: Dict[str, Any],
        cash_flow: Dict[str, Any],
        industry: Dict[str, Any],
        promoter: Dict[str, Any],
        external_intelligence: Dict[str, Any]
    ) -> str:
        """
        Build a formatted table of risk factor scores.
        
        Args:
            financial_health: Financial health factor data
            cash_flow: Cash flow factor data
            industry: Industry factor data
            promoter: Promoter factor data
            external_intelligence: External intelligence factor data
        
        Returns:
            Markdown table of risk factors
        
        Requirements: 7.3
        """
        table = "| Risk Factor | Score | Weight | Key Findings |\n"
        table += "|-------------|-------|--------|-------------|\n"
        
        factors = [
            ('Financial Health', financial_health, 0.35),
            ('Cash Flow', cash_flow, 0.25),
            ('Industry', industry, 0.15),
            ('Promoter', promoter, 0.15),
            ('External Intelligence', external_intelligence, 0.10)
        ]
        
        for factor_name, factor_data, weight in factors:
            score = factor_data.get('score', 0)
            key_findings = factor_data.get('key_findings', [])
            
            # Format score with color indicator
            if score >= 70:
                score_str = f"✓ {score:.1f}"
            elif score >= 50:
                score_str = f"~ {score:.1f}"
            else:
                score_str = f"✗ {score:.1f}"
            
            # Format weight
            weight_str = f"{weight * 100:.0f}%"
            
            # Format key findings (limit to 2)
            findings_str = "; ".join(key_findings[:2]) if key_findings else "N/A"
            if len(findings_str) > 80:
                findings_str = findings_str[:77] + "..."
            
            table += f"| {factor_name} | {score_str} | {weight_str} | {findings_str} |\n"
        
        return table

    async def export_to_pdf(self, cam: CAM) -> bytes:
        """
        Export CAM to PDF format using WeasyPrint.
        
        Args:
            cam: CAM object to export
        
        Returns:
            PDF file content as bytes
        
        Raises:
            ValueError: If CAM content is invalid or export fails
        
        Requirements: 7.4
        """
        try:
            # Convert markdown to HTML
            html_content = markdown2.markdown(
                cam.content,
                extras=['tables', 'fenced-code-blocks', 'header-ids']
            )
            
            # Build complete HTML document with CSS styling
            full_html = self._build_html_document(html_content, cam)
            
            # Generate PDF using WeasyPrint (updated API)
            html_doc = HTML(string=full_html)
            pdf_bytes = html_doc.write_pdf()
            
            return pdf_bytes
        
        except Exception as e:
            raise ValueError(f"Failed to export CAM to PDF: {str(e)}")

    async def export_to_word(self, cam: CAM) -> bytes:
        """
        Export CAM to Word (DOCX) format using python-docx.
        
        Args:
            cam: CAM object to export
        
        Returns:
            Word document content as bytes
        
        Raises:
            ValueError: If CAM content is invalid or export fails
        
        Requirements: 7.4
        """
        try:
            # Create new Word document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = "Credit Appraisal Memo"
            doc.core_properties.created = cam.generated_at
            
            # Add content to document
            self._add_content_to_word_doc(doc, cam)
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return doc_bytes.getvalue()
        
        except Exception as e:
            raise ValueError(f"Failed to export CAM to Word: {str(e)}")

    def validate_export_format(self, format: str) -> bool:
        """
        Validate that the requested export format is supported.
        
        Args:
            format: Export format ('pdf' or 'word')
        
        Returns:
            True if format is valid, False otherwise
        
        Requirements: 7.4
        """
        return format.lower() in ['pdf', 'word', 'docx']

    def _build_html_document(self, html_content: str, cam: CAM) -> str:
        """
        Build a complete HTML document with CSS styling for PDF export.
        
        Args:
            html_content: Converted markdown content as HTML
            cam: CAM object for metadata
        
        Returns:
            Complete HTML document as string
        """
        css_styles = """
        <style>
            @page {
                size: A4;
                margin: 2cm;
            }
            body {
                font-family: 'Helvetica', 'Arial', sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #333;
            }
            h1 {
                color: #1a1a1a;
                font-size: 24pt;
                margin-top: 0;
                margin-bottom: 20pt;
                border-bottom: 2pt solid #2563eb;
                padding-bottom: 10pt;
            }
            h2 {
                color: #2563eb;
                font-size: 18pt;
                margin-top: 20pt;
                margin-bottom: 12pt;
                border-bottom: 1pt solid #e5e7eb;
                padding-bottom: 6pt;
            }
            h3 {
                color: #1e40af;
                font-size: 14pt;
                margin-top: 16pt;
                margin-bottom: 8pt;
            }
            p {
                margin-bottom: 10pt;
                text-align: justify;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                margin: 15pt 0;
                font-size: 10pt;
            }
            th {
                background-color: #2563eb;
                color: white;
                padding: 8pt;
                text-align: left;
                font-weight: bold;
            }
            td {
                padding: 6pt 8pt;
                border-bottom: 1pt solid #e5e7eb;
            }
            tr:nth-child(even) {
                background-color: #f9fafb;
            }
            hr {
                border: none;
                border-top: 1pt solid #d1d5db;
                margin: 20pt 0;
            }
            .metadata {
                color: #6b7280;
                font-size: 9pt;
                margin-top: 10pt;
            }
            strong {
                color: #1a1a1a;
            }
        </style>
        """
        
        html_doc = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Credit Appraisal Memo</title>
            {css_styles}
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        return html_doc

    def _add_content_to_word_doc(self, doc: Document, cam: CAM) -> None:
        """
        Add CAM content to a Word document with proper formatting.
        
        Args:
            doc: python-docx Document object
            cam: CAM object with content to add
        """
        # Parse markdown content into sections
        lines = cam.content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Handle headers
            if line.startswith('# '):
                # H1 - Main title
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('## '):
                # H2 - Section headers
                heading = doc.add_heading(line[3:], level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('### '):
                # H3 - Subsection headers
                heading = doc.add_heading(line[4:], level=3)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Handle horizontal rules
            elif line.startswith('---'):
                doc.add_paragraph('_' * 50)
            
            # Handle tables (basic markdown table parsing)
            elif line.startswith('|'):
                # Skip table processing for now - tables are complex
                # In production, you'd want proper markdown table parsing
                doc.add_paragraph(line)
            
            # Handle bold text with **
            elif '**' in line:
                para = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        # Normal text
                        para.add_run(part)
                    else:
                        # Bold text
                        para.add_run(part).bold = True
            
            # Handle list items
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            
            # Regular paragraphs
            else:
                doc.add_paragraph(line)

    # Fallback methods for when AI generation fails
    
    def _generate_fallback_executive_summary(
        self,
        company_name: str,
        loan_amount: float,
        overall_score: float,
        recommendation: str
    ) -> str:
        """Generate rule-based executive summary as fallback."""
        rec_text = recommendation.replace('_', ' ').title()
        
        if overall_score >= 70:
            assessment = "strong creditworthiness"
        elif overall_score >= 50:
            assessment = "moderate creditworthiness"
        else:
            assessment = "weak creditworthiness"
        
        return f"""This Credit Appraisal Memo evaluates {company_name}'s application for a loan of ${loan_amount:,.2f}. Based on comprehensive financial analysis, multi-source intelligence gathering, and risk assessment, the company has been assigned a credit score of {overall_score:.1f}/100, indicating {assessment}.

The recommendation is to **{rec_text}** the loan application. This assessment is based on analysis of financial health, cash flow adequacy, industry position, management quality, and external intelligence.

Key factors considered include the company's financial ratios, historical performance trends, future growth projections, and overall risk profile. Detailed analysis and supporting rationale are provided in the sections below.
"""
    
    def _generate_fallback_company_overview(
        self,
        company_name: str,
        loan_purpose: str,
        research_data: Dict[str, Any]
    ) -> str:
        """Generate rule-based company overview as fallback."""
        return f"""### Company Background

{company_name} has applied for financing to support {loan_purpose}. The company operates in a competitive market environment and has been evaluated based on available business intelligence and market research.

### Management and Promoters

Management analysis has been conducted to assess the experience, track record, and reputation of the company's leadership team. This evaluation considers the promoters' background, previous business ventures, and industry expertise.

### Industry Position

The company's position within its industry sector has been analyzed, including competitive dynamics, market trends, and growth potential. Industry-specific risks and opportunities have been identified and factored into the credit assessment.

### Loan Purpose

The requested loan of will be used for {loan_purpose}. The viability and appropriateness of this use of funds has been evaluated as part of the overall credit analysis.
"""
    
    def _generate_fallback_financial_narrative(
        self,
        ratios: Dict[str, Any],
        trends: Dict[str, Any],
        benchmarks: Dict[str, Any]
    ) -> str:
        """Generate rule-based financial narrative as fallback."""
        narrative = "The financial analysis encompasses a comprehensive review of the company's financial position, including liquidity, leverage, profitability, and efficiency metrics.\n\n"
        
        # Assess liquidity
        current_ratio = ratios.get('current_ratio', {}).get('value', 0)
        if current_ratio >= 1.5:
            narrative += "The company demonstrates strong liquidity with a healthy current ratio, indicating adequate short-term financial flexibility. "
        elif current_ratio >= 1.0:
            narrative += "The company maintains acceptable liquidity levels, though there is limited cushion for short-term obligations. "
        else:
            narrative += "Liquidity is a concern, with the current ratio indicating potential challenges in meeting short-term obligations. "
        
        # Assess leverage
        debt_to_equity = ratios.get('debt_to_equity', {}).get('value', 0)
        if debt_to_equity <= 1.0:
            narrative += "The capital structure shows moderate leverage with a balanced debt-to-equity ratio. "
        else:
            narrative += "The company carries significant leverage, which increases financial risk. "
        
        # Assess profitability
        roe = ratios.get('roe', {}).get('value', 0)
        if roe >= 0.15:
            narrative += "Profitability metrics are strong, with healthy returns on equity indicating efficient use of shareholder capital.\n\n"
        elif roe >= 0.05:
            narrative += "Profitability is moderate, with returns on equity meeting acceptable standards.\n\n"
        else:
            narrative += "Profitability is weak, with low returns on equity raising concerns about operational efficiency.\n\n"
        
        # Assess trends
        if trends:
            narrative += "Historical trend analysis reveals "
            positive_trends = sum(1 for t in trends.values() if t.get('trend_direction') in ['increasing', 'stable'])
            if positive_trends > len(trends) / 2:
                narrative += "generally positive momentum in key financial metrics, supporting the company's growth trajectory. "
            else:
                narrative += "concerning patterns in several key metrics, indicating potential operational challenges. "
        
        narrative += "These factors have been considered in the overall credit risk assessment."
        
        return narrative

    def _generate_fallback_risk_narrative(
        self,
        overall_score: float,
        recommendation: str,
        financial_health: Dict[str, Any],
        cash_flow: Dict[str, Any]
    ) -> str:
        """Generate rule-based risk narrative as fallback."""
        rec_text = recommendation.replace('_', ' ').title()
        
        narrative = f"The comprehensive risk assessment has resulted in an overall credit score of {overall_score:.1f}/100, leading to a recommendation to **{rec_text}** the loan application.\n\n"
        
        # Assess overall risk level
        if overall_score >= 70:
            narrative += "This score indicates strong creditworthiness with limited credit risk. The company demonstrates solid financial fundamentals, adequate cash flow generation, and favorable market positioning. "
        elif overall_score >= 50:
            narrative += "This score indicates moderate creditworthiness with manageable credit risk. While the company shows acceptable performance in several areas, there are some concerns that warrant monitoring. "
        else:
            narrative += "This score indicates weak creditworthiness with elevated credit risk. The company exhibits significant weaknesses across multiple risk factors that raise substantial concerns. "
        
        # Highlight key factors
        fh_score = financial_health.get('score', 0)
        cf_score = cash_flow.get('score', 0)
        
        narrative += f"\n\nFinancial health scored {fh_score:.1f}/100, reflecting the company's balance sheet strength and ratio performance. "
        narrative += f"Cash flow scored {cf_score:.1f}/100, indicating the company's ability to generate and sustain adequate cash flows. "
        
        # Risk conclusion
        if overall_score >= 70:
            narrative += "\n\nThe overall risk profile is favorable, with strengths outweighing weaknesses. The company is well-positioned to service the requested debt."
        elif overall_score >= 50:
            narrative += "\n\nThe overall risk profile is mixed, with both strengths and weaknesses present. Conditional approval with appropriate covenants and monitoring is recommended."
        else:
            narrative += "\n\nThe overall risk profile is unfavorable, with significant weaknesses that outweigh any strengths. The credit risk is too high to support approval at this time."
        
        return narrative
    
    def _generate_fallback_recommendation(
        self,
        recommendation: str,
        overall_score: float,
        loan_amount: float
    ) -> str:
        """Generate rule-based recommendation as fallback."""
        rec_text = recommendation.replace('_', ' ').upper()
        
        content = f"Based on the comprehensive credit analysis, the recommendation is to **{rec_text}** the loan application for ${loan_amount:,.2f}.\n\n"
        
        if recommendation == 'approve':
            content += f"""**Rationale:** The company has demonstrated strong creditworthiness with a credit score of {overall_score:.1f}/100. The financial analysis reveals solid fundamentals, adequate cash flow generation, and favorable risk factors. The company is well-positioned to service the requested debt.

**Recommended Terms:**
- Standard loan terms and conditions
- Regular financial reporting requirements
- Annual credit review

**Monitoring:** Standard post-approval monitoring should be implemented to track financial performance and covenant compliance.
"""
        elif recommendation == 'approve_with_conditions':
            content += f"""**Rationale:** The company has demonstrated moderate creditworthiness with a credit score of {overall_score:.1f}/100. While the overall assessment is acceptable, certain weaknesses have been identified that require mitigation through conditions and enhanced monitoring.

**Recommended Conditions:**
- Enhanced financial reporting (quarterly)
- Maintenance of minimum financial covenants (current ratio, debt service coverage)
- Restrictions on additional debt without lender approval
- Regular business updates and management meetings

**Monitoring:** Enhanced post-approval monitoring should be implemented with quarterly reviews and covenant testing.
"""
        else:  # reject
            content += f"""**Rationale:** The company has demonstrated weak creditworthiness with a credit score of {overall_score:.1f}/100. The analysis has identified significant concerns across multiple risk factors that make the credit risk unacceptable at this time.

**Key Concerns:**
- Weak financial fundamentals and ratio performance
- Inadequate cash flow generation
- Elevated leverage or liquidity concerns
- Unfavorable industry or management factors

**Recommendation:** The loan application should be declined. The company may reapply once the identified concerns have been adequately addressed and financial performance has improved.
"""
        
        return content
