"""
Document Processing Service

This module provides document text extraction functionality for various file formats
including PDF, DOCX, Excel, CSV, and images.

Requirements: 1.1, 1.2, 1.3, 15.4
"""

import io
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from enum import Enum

import PyPDF2
import openpyxl
import pandas as pd
from docx import Document
from PIL import Image


class FileType(Enum):
    """Supported file types for document processing"""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    XLS = "xls"
    CSV = "csv"
    JPG = "jpg"
    JPEG = "jpeg"
    PNG = "png"


class DocumentProcessingError(Exception):
    """Base exception for document processing errors"""
    pass


class UnsupportedFileTypeError(DocumentProcessingError):
    """Raised when file type is not supported"""
    pass


class CorruptedFileError(DocumentProcessingError):
    """Raised when file is corrupted or cannot be processed"""
    pass


class FileSizeExceededError(DocumentProcessingError):
    """Raised when file size exceeds maximum allowed"""
    pass


class DocumentProcessor:
    """
    Handles document text extraction for multiple file formats.
    
    Supports:
    - PDF documents
    - Word documents (DOCX)
    - Excel spreadsheets (XLSX, XLS)
    - CSV files
    - Images (JPG, JPEG, PNG)
    
    Requirements: 1.1, 1.2, 1.3, 15.4
    """
    
    # Maximum file size: 50MB
    MAX_FILE_SIZE = 50 * 1024 * 1024
    
    # Supported MIME types
    SUPPORTED_MIME_TYPES = {
        'application/pdf': FileType.PDF,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': FileType.DOCX,
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': FileType.XLSX,
        'application/vnd.ms-excel': FileType.XLS,
        'text/csv': FileType.CSV,
        'image/jpeg': FileType.JPEG,
        'image/jpg': FileType.JPG,
        'image/png': FileType.PNG,
    }
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        '.pdf': FileType.PDF,
        '.docx': FileType.DOCX,
        '.xlsx': FileType.XLSX,
        '.xls': FileType.XLS,
        '.csv': FileType.CSV,
        '.jpg': FileType.JPG,
        '.jpeg': FileType.JPEG,
        '.png': FileType.PNG,
    }
    
    def detect_file_type(self, filename: str, content: Optional[bytes] = None) -> FileType:
        """
        Detect file type from filename and optionally content.
        
        Args:
            filename: Name of the file
            content: Optional file content for MIME type detection
            
        Returns:
            FileType enum value
            
        Raises:
            UnsupportedFileTypeError: If file type is not supported
        """
        # Try extension-based detection first
        file_path = Path(filename)
        extension = file_path.suffix.lower()
        
        if extension in self.SUPPORTED_EXTENSIONS:
            return self.SUPPORTED_EXTENSIONS[extension]
        
        # Try MIME type detection if content is provided
        if content:
            mime_type, _ = mimetypes.guess_type(filename)
            if mime_type and mime_type in self.SUPPORTED_MIME_TYPES:
                return self.SUPPORTED_MIME_TYPES[mime_type]
        
        raise UnsupportedFileTypeError(
            f"Unsupported file type: {extension}. "
            f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS.keys())}"
        )
    
    def validate_file(self, filename: str, content: bytes) -> Tuple[FileType, bool]:
        """
        Validate file format and size.
        
        Args:
            filename: Name of the file
            content: File content as bytes
            
        Returns:
            Tuple of (FileType, is_valid)
            
        Raises:
            FileSizeExceededError: If file size exceeds maximum
            UnsupportedFileTypeError: If file type is not supported
        """
        # Check file size
        file_size = len(content)
        if file_size > self.MAX_FILE_SIZE:
            raise FileSizeExceededError(
                f"File size ({file_size} bytes) exceeds maximum allowed "
                f"({self.MAX_FILE_SIZE} bytes)"
            )
        
        # Detect and validate file type
        file_type = self.detect_file_type(filename, content)
        
        return file_type, True
    
    def extract_text(self, filename: str, content: bytes) -> str:
        """
        Extract text from document based on file type.
        
        Args:
            filename: Name of the file
            content: File content as bytes
            
        Returns:
            Extracted text content
            
        Raises:
            UnsupportedFileTypeError: If file type is not supported
            CorruptedFileError: If file is corrupted or cannot be processed
            FileSizeExceededError: If file size exceeds maximum
        """
        # Validate file first
        file_type, _ = self.validate_file(filename, content)
        
        try:
            # Route to appropriate extraction method
            if file_type == FileType.PDF:
                return self._extract_from_pdf(content)
            elif file_type == FileType.DOCX:
                return self._extract_from_docx(content)
            elif file_type in [FileType.XLSX, FileType.XLS]:
                return self._extract_from_excel(content)
            elif file_type == FileType.CSV:
                return self._extract_from_csv(content)
            elif file_type in [FileType.JPG, FileType.JPEG, FileType.PNG]:
                return self._extract_from_image(content)
            else:
                raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")
                
        except (UnsupportedFileTypeError, FileSizeExceededError):
            # Re-raise validation errors
            raise
        except Exception as e:
            # Wrap other errors as corrupted file errors
            raise CorruptedFileError(
                f"Failed to process file '{filename}': {str(e)}"
            ) from e
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF document.
        
        Args:
            content: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"[Page {page_num}]\n{page_text}")
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise CorruptedFileError(f"Failed to extract text from PDF: {str(e)}") from e
    
    def _extract_from_docx(self, content: bytes) -> str:
        """
        Extract text from Word document.
        
        Args:
            content: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = self._extract_table_from_docx(table)
                if table_text:
                    text_parts.append(table_text)
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise CorruptedFileError(f"Failed to extract text from DOCX: {str(e)}") from e
    
    def _extract_table_from_docx(self, table) -> str:
        """Extract text from a Word table"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Only include non-empty rows
                rows.append(" | ".join(cells))
        return "\n".join(rows)
    
    def _extract_from_excel(self, content: bytes) -> str:
        """
        Extract text from Excel spreadsheet.
        
        Args:
            content: Excel file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            excel_file = io.BytesIO(content)
            workbook = openpyxl.load_workbook(excel_file, data_only=True)
            
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"[Sheet: {sheet_name}]")
                
                # Extract data from sheet
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    # Convert row values to strings, filtering out None
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_values):  # Only include non-empty rows
                        rows.append(" | ".join(row_values))
                
                if rows:
                    text_parts.append("\n".join(rows))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise CorruptedFileError(f"Failed to extract text from Excel: {str(e)}") from e
    
    def _extract_from_csv(self, content: bytes) -> str:
        """
        Extract text from CSV file.
        
        Args:
            content: CSV file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            csv_file = io.BytesIO(content)
            df = pd.read_csv(csv_file)
            
            # Convert DataFrame to text representation
            text_parts = []
            
            # Add column headers
            headers = " | ".join(df.columns)
            text_parts.append(headers)
            text_parts.append("-" * len(headers))
            
            # Add rows
            for _, row in df.iterrows():
                row_text = " | ".join(str(val) for val in row.values)
                text_parts.append(row_text)
            
            return "\n".join(text_parts)
            
        except Exception as e:
            raise CorruptedFileError(f"Failed to extract text from CSV: {str(e)}") from e
    
    def _extract_from_image(self, content: bytes) -> str:
        """
        Extract text from image file.
        
        Note: This is a placeholder implementation. For production use,
        OCR functionality (e.g., using Tesseract or cloud OCR services)
        should be integrated here.
        
        Args:
            content: Image file content as bytes
            
        Returns:
            Extracted text (currently returns image metadata)
        """
        try:
            image_file = io.BytesIO(content)
            image = Image.open(image_file)
            
            # For now, return image metadata
            # In production, integrate OCR service here
            metadata = {
                "format": image.format,
                "mode": image.mode,
                "size": f"{image.width}x{image.height}",
            }
            
            text = "[Image File]\n"
            text += "\n".join(f"{key}: {value}" for key, value in metadata.items())
            text += "\n\nNote: OCR text extraction not yet implemented for images."
            
            return text
            
        except Exception as e:
            raise CorruptedFileError(f"Failed to process image: {str(e)}") from e
    
    def extract_with_metadata(
        self, 
        filename: str, 
        content: bytes
    ) -> Dict[str, any]:
        """
        Extract text along with metadata about the extraction.
        
        Args:
            filename: Name of the file
            content: File content as bytes
            
        Returns:
            Dictionary containing:
                - text: Extracted text
                - file_type: Detected file type
                - file_size: Size in bytes
                - page_count: Number of pages (for PDFs)
                - sheet_count: Number of sheets (for Excel)
        """
        file_type, _ = self.validate_file(filename, content)
        text = self.extract_text(filename, content)
        
        metadata = {
            "text": text,
            "file_type": file_type.value,
            "file_size": len(content),
            "filename": filename,
        }
        
        # Add format-specific metadata
        if file_type == FileType.PDF:
            try:
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                metadata["page_count"] = len(pdf_reader.pages)
            except:
                pass
        
        elif file_type in [FileType.XLSX, FileType.XLS]:
            try:
                excel_file = io.BytesIO(content)
                workbook = openpyxl.load_workbook(excel_file, data_only=True)
                metadata["sheet_count"] = len(workbook.sheetnames)
            except:
                pass
        
        return metadata
